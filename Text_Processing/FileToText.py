import re

import docx

from pptx import Presentation

from bs4 import BeautifulSoup

from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFEncryptionError, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser

from Exceptions.OwnExceptions import ExceptionReadDataToText
from Text_Processing.IFileToText import IFileToText


class FileToText(IFileToText):
    def fileToText(self, filePath, datatype):
        # set the datatype on lowercase
        datatype = datatype.lower()
        # check if the datatype is supported
        # if not return an empty string
        if datatype == 'pdf':
            text = self._readPDFAsLectureScript(filePath)
        elif datatype == 'pptx':
            text = self._readPPTX(filePath)
        elif datatype == 'docx':
            text = self._readDOCX(filePath)
        elif datatype == 'html':
            text = self._htmlToText(filePath)
        else:
            text = ""
        return text

    def _readDOCX(self, docxPath: str) -> str:
        """readDOCX read the text from a docx-file.

            :param docxPath: The path to the file
            :return Return the text from the file
        """
        doc = docx.Document(docxPath)
        text = ""
        # for each paragraph add the text to the result text
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def _readPPTX(self, pptxPath: str) -> str:
        """readPPTX read the text from a pptx-file.
            :param pptxPath: The path to the file
            :return Return the text from the file
        """
        # regeluar expression to find all required text in xml-blocks in pptx
        regexText = r"\<a:t\>(.*)\<\/a:t\>"

        # open the presentation
        prs = Presentation(pptxPath)
        # get all slides from the presentation
        list_slide = prs.slides

        text = ""
        # iterate through every slide
        for slide in list_slide:
            # get for the slide all blocks where text could be stored
            blocks = self._getTextBlocks(slide._element.xml)
            # iterate through all blocks
            for block in blocks:
                # set has_text on false
                has_text = False
                # try to find in block the regular expression for the text
                for tempText in re.findall(regexText, block):
                    # add the found text to the result text
                    text += tempText
                    has_text = True
                # if a text was found in this block write a ". " to the end of the result text
                if (has_text):
                    text += ". "
        return text

    def _getTextBlocks(self, pptxxml: str) -> list:
        """getTextBlocks return a list of start- and entpoints which could contain text in a pptx slide.

                :param pptxxml: the xml of a pptx slide
                :return Returns a list of start- and entpoints from blocks
                :raise ExceptionReadDataToText: if the length of start- and endpoints do not match
        """
        startPos = []
        endPos = []
        blocks = []
        # regular expression for the searched strings
        # which represent the start and end of a xml-block,
        # which can contain text
        regexBlockBeginn = r"\<a:p\>"
        regexBlockEnd = r"\<\/a:p\>"
        # find all occurrence of the regex
        for m in re.finditer(regexBlockBeginn, pptxxml):
            startPos.append(m.start())
        for m in re.finditer(regexBlockEnd, pptxxml):
            endPos.append(m.start())
        # if the number of start- and endpoints do not match
        # raise ExceptionReadDataToText
        if (len(startPos) != len(endPos)):
            raise ExceptionReadDataToText("PPTX Blockanfang und Ende passt nicht")
        # create the blocks
        for j in range(len(startPos)):
            blocks.append(pptxxml[startPos[j] + 1:endPos[j] - 1])
        return blocks

    def _readPDFAsLectureScript(self, pdfPath: str) -> str:
        """readPDFAsLectureScript read the text from a pdf-file. The special on this script is, that
        after every line break a ". " will be inserted. Because we assume, that we read a lecture script and
        a lecture script has mostly only bullet points and no real "." in the text.

            :param pdfPath: The path to the file
            :return The text from the file
            :raise ExceptionReadDataToText: if the pdf is not extractable or with a password saved. Also this Exception will be trhwon if an problem occure.
        """
        try:
            # Create the document model from the file
            file = open(pdfPath, 'rb')
            parser = PDFParser(file)
            document = PDFDocument(parser)
            # Try to parse the document
            if not document.is_extractable:
                raise ExceptionReadDataToText('Es ist nicht erlaubt die PDF zu parsen')
            # Create a PDF resource manager object
            # that stores shared resources.
            rsrcmgr = PDFResourceManager()
            # Create a buffer for the parsed text
            retstr = StringIO()
            # Spacing parameters for parsing
            laparams = LAParams()
            codec = 'utf-8'
            # Create a PDF device object
            device = TextConverter(rsrcmgr, retstr,
                                   codec=codec,
                                   laparams=laparams)
            # Create a PDF interpreter object
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            # Process each page contained in the document.
            for page in PDFPage.create_pages(document):
                interpreter.process_page(page)
            # replace the occurrence of multiple points with "."
            # replace the occurrence of multiple new lines with a ". \n"
            text = re.sub(r' {0,}\.+ \n', '. \n', re.sub('\n+', '. \n', retstr.getvalue()))
            # (cid:int) is from the pdfToText characters, which the framework can not translate
            text = re.sub(r'\(cid:\d+\)', '', text)
            text = re.sub(r'\f', '', text)
            return text
        except PDFEncryptionError as e:
            raise ExceptionReadDataToText('Die Datei ist verschluesselt und kann nicht geoeffnet werden')
        finally:
            # close the file
            file.close()

    def _htmlToText(self, htmlPath: str) -> str:
        """htmlToText return the text from the html-file which is stored in the htmlPath

                :param htmlPath: the path to the html file
                :return Returns the text from the html file (str)
                :raise ExceptionReadDataToText if the coding cant be handeled
        """
        try:
            # open the file with utf 8 encoding
            file = open(htmlPath, encoding="utf8")
            htmltext = ""
            # fill the htmltext
            for line in file:
                htmltext += line
        except UnicodeDecodeError as e:
            # when an Unicode Decode Error occurs, read the text without encoding
            try:
                file = open(htmlPath)
                htmltext = ""
                for line in file:
                    htmltext += line
            except UnicodeDecodeError as e:
                raise ExceptionReadDataToText('Die Formatierung kann nicht behandelt werden.')
        finally:
            # if the file is not closed than close the file
            if not file.closed:
                file.close()

        # parse the htmltext with BeautifulSoup
        parsed_article = BeautifulSoup(htmltext, 'lxml')
        # find all paragraphs from headers to normal text, to bullet points
        paragraphs = parsed_article.find_all(re.compile('(^h[1-6]|^p|^li)'))
        article_text = ""

        # iterate through every paragraph
        for p in paragraphs:
            # if the paragraphs is a header or a bullet point
            # add to article_text the text of the paragraph and add to the end ". ".
            # Because this are own sentences
            if p.name in ['li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                article_text += p.text + ". "
            # otherwise add only the text
            else:
                article_text += p.text
        # replace any multiple whitespace areas with one " "
        article_text = re.sub(r'\s+', ' ', article_text)
        return article_text
